#!/usr/bin/env python
# coding: utf-8

# In[ ]:


get_ipython().set_next_input('Q1). In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened');get_ipython().run_line_magic('pinfo', 'opened')


# In[ ]:


Answer:-
PyPDF2.PdfFileReader() needs to be opened in read-binary mode by passing 'rb' as the second argument to open(). Likewise,
the File object passed to PyPDF2.PdfFileWriter() needs to be opened in write-binary mode with 'wb'.


# In[ ]:


Q2)From a PdfFileReader object, how do you get a Page object for page 5?


# In[ ]:


Answer:-We can get a Page object by calling the getPage(5).


# In[ ]:


get_ipython().set_next_input('Q3)What PdfFileReader variable stores the number of pages in the PDF document');get_ipython().run_line_magic('pinfo', 'document')


# In[ ]:


Answer:-The total number of pages in the document is stored in the numPages attribute of a PdfFileReader object


# In[ ]:


Q4) If a PdfFileReader object’s PDF is encrypted with the password swordfish,
get_ipython().set_next_input('what must you do before you can obtain Page objects from it');get_ipython().run_line_magic('pinfo', 'it')


# In[ ]:


Amswer:-
    Before we obtain the page object, the pdf has to be decrypted by calling .decrypt('swordfish')


# In[ ]:


get_ipython().set_next_input('Q5)hat is the difference between a Run object and a Paragraph object');get_ipython().run_line_magic('pinfo', 'object')


# In[ ]:


Paragraph Object : A document contains multiple paragraphs. A paragraph begins on a new line and contains multiple
runs. The Document object contains a list of Paragraph objects for the paragraphs in the document.
(A new paragraph begins whenever the user presses ENTER or RETURN while typing in a Word document.)

Run Objects : Runs are contiguous groups of characters within a paragraph with the same style


# In[ ]:


get_ipython().set_next_input('Q7)How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc');get_ipython().run_line_magic('pinfo', 'doc')


# In[ ]:


#!pip install python-docx
import docx
doc = docx.Document('abc.docx')
doc.paragraphs
#By using doc.paragraphs


# In[ ]:


get_ipython().set_next_input('Q8)What type of object has bold, underline, italic, strike, and outline variables');get_ipython().run_line_magic('pinfo', 'variables')


# In[ ]:


Answer:- A Run object has bold, underline,italic,strike and outline variables


# In[ ]:


get_ipython().set_next_input('Q9). What is the difference between False, True, and None for the bold variable');get_ipython().run_line_magic('pinfo', 'variable')


# In[ ]:


Answer:- Runs can be further styled using text attributes. Each attribute can be set to one of three values:
True (the attribute is always enabled, no matter what other styles are applied to the run),
False (the attribute is always disabled),
None (defaults to whatever the run’s style is set to)
True always makes the Run object bolded and False makes it always not bolded, no matter what the style’s bold setting is. 
None will make the Run object just use the style’s bold setting


# In[ ]:


get_ipython().set_next_input('Q10). How do you create a Document object for a new Word document');get_ipython().run_line_magic('pinfo', 'document')


# In[ ]:


Answer:-
    By Calling the docx.Document() function


# In[ ]:


Q11)How do you add a paragraph with the text 'Hello, there!' to a Document object stored in a variable named doc


# In[ ]:


import docx
doc = docx.Document()

doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


# In[ ]:


get_ipython().set_next_input('Q12)What integers represent the levels of headings available in Word documents');get_ipython().run_line_magic('pinfo', 'documents')


# In[ ]:


integer from 0 to 4
The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. 
The integer 0 makes the heading the Title style, which is used for the top of the document. 
Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading

